import io

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _set_cell_bg(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_col_widths(table, widths_cm: list[float]) -> None:
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = Cm(widths_cm[i])


def export_plan_to_docx(plan: list[dict]) -> bytes:
    doc = Document()

    # Page margins
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Контент-план Петрострой")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x1A, 0x56, 0x7A)

    # Month subtitle
    if plan:
        first = plan[0]
        last = plan[-1]
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_run = subtitle.add_run(
            f"{first['date_str']} — {last['date_str']}"
        )
        sub_run.font.size = Pt(11)
        sub_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()

    # Table: №, Дата, Рубрика, Тема, Примечания
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"

    headers = ["№", "Дата", "Рубрика", "Тема поста", "Примечания"]
    col_widths = [1.0, 3.0, 3.5, 8.0, 3.5]

    hdr_row = table.rows[0]
    for i, (cell, header) in enumerate(zip(hdr_row.cells, headers)):
        _set_cell_bg(cell, "1A567A")
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(header)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for item in plan:
        row = table.add_row()
        bg = "EAF3F8" if item["num"] % 2 == 0 else "FFFFFF"
        values = [
            str(item["num"]),
            item["date_str"],
            f"{item['rubric_emoji']} {item['rubric_name']}",
            item["topic"],
            "",
        ]
        for i, (cell, val) in enumerate(zip(row.cells, values)):
            _set_cell_bg(cell, bg)
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER if i < 3 else WD_ALIGN_PARAGRAPH.LEFT
            run = para.add_run(val)
            run.font.size = Pt(10)

    _set_col_widths(table, col_widths)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
